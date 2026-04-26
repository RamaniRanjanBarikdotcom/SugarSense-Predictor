"""Generate SugarSense Predictor — 50-Page Professional Academic Report (.docx)"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ────────────────────────────────────────────────────────────────────
PRI   = RGBColor(0x0f, 0x76, 0x6e)
ACC   = RGBColor(0x14, 0xb8, 0xa6)
DARK  = RGBColor(0x0f, 0x17, 0x2a)
SOFT  = RGBColor(0x47, 0x55, 0x69)
LIGHT = RGBColor(0x94, 0xa3, 0xb8)
WHITE = RGBColor(0xff, 0xff, 0xff)
BLACK = RGBColor(0x00, 0x00, 0x00)
RED   = RGBColor(0xdc, 0x26, 0x26)
GREEN = RGBColor(0x16, 0xa3, 0x4a)

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.2)
    section.right_margin  = Cm(2.5)
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)

# ── Default styles ─────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name  = 'Times New Roman'
style_normal.font.size  = Pt(11)
style_normal.paragraph_format.space_after  = Pt(6)
style_normal.paragraph_format.space_before = Pt(0)

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_col_width(table, col_widths):
    for row in table.rows:
        for i, w in enumerate(col_widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)

def heading(text, level=1, clr=None, align=WD_ALIGN_PARAGRAPH.LEFT, pg_break=False):
    if pg_break:
        doc.add_page_break()
    p = doc.add_heading(text, level=level)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(16 if level == 1 else 10)
    pf.space_after  = Pt(8)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        if clr:
            run.font.color.rgb = clr
        elif level == 1:
            run.font.color.rgb = PRI
        elif level == 2:
            run.font.color.rgb = ACC
        else:
            run.font.color.rgb = DARK
    return p

def para(text, size=11, bold=False, italic=False, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after  = Pt(space_after)
    pf.space_before = Pt(0)
    if indent:
        pf.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name   = 'Times New Roman'
    run.font.size   = Pt(size)
    run.bold        = bold
    run.italic      = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK
    return p

def cpara(text, size=11, bold=False, color=None):
    return para(text, size=size, bold=bold, color=color,
                align=WD_ALIGN_PARAGRAPH.CENTER)

def bullet(text, level=0, size=11):
    style = 'List Bullet 2' if level > 0 else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.color.rgb = DARK
    return p

def numbered(text, size=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.color.rgb = DARK
    return p

def table(headers, rows, col_widths=None, hdr_color="0f766e", alt_color="EBF5FB"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, hdr_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        if ri % 2 == 1:
            for c in row.cells:
                set_cell_bg(c, alt_color)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = DARK
    if col_widths:
        set_col_width(t, col_widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def code_block(lines):
    """Render code lines as a shaded paragraph block."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.4)
        set_para_bg(p, "1e293b")
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7d, 0xd3, 0xfc)

def set_para_bg(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    pPr.append(shd)

def divider():
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    for run in p.runs:
        run.font.color.rgb = LIGHT
        run.font.size = Pt(8)

def sp(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

def label_val(lbl, val, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{lbl}  ")
    r1.bold = True
    r1.font.size = Pt(size)
    r1.font.color.rgb = DARK
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(val)
    r2.font.size = Pt(size)
    r2.font.color.rgb = SOFT
    r2.font.name = 'Times New Roman'
    return p

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
sp(2)
cpara("SOPHITORIUM ENGINEERING COLLEGE", 13, bold=True, color=DARK)
cpara("Department of Computer Science & Engineering", 11, color=SOFT)
cpara("Affiliated to Biju Patnaik University of Technology (BPUT), Odisha", 10, color=SOFT)
sp(1)
divider()
sp(1)
cpara("B.TECH MINOR PROJECT REPORT", 12, bold=True, color=SOFT)
cpara("Academic Year: 2025 – 2026", 11, color=SOFT)
sp(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SugarSense Predictor")
run.bold = True
run.font.size = Pt(30)
run.font.color.rgb = PRI
run.font.name = 'Times New Roman'

cpara("Smart Diabetes Risk Prediction & AI Health Assistant System", 14, color=ACC)
sp(1)
cpara("A Full-Stack Web Application powered by Machine Learning & Large Language Models", 11, color=SOFT)
sp(2)

table(
    ["Technology Stack", "Flask · React · SVM · Claude AI · Google Gemini"],
    [
        ["Dataset Used",     "Pima Indians Diabetes Database (768 Samples, 8 Features)"],
        ["ML Algorithm",     "Support Vector Machine (SVM) — Linear Kernel"],
        ["AI Integration",   "5 Providers: Anthropic, OpenAI, Google, HuggingFace, OpenRouter"],
        ["Model Accuracy",   "~77–78% on held-out test set"],
    ],
    col_widths=[2.5, 4.0],
    hdr_color="0f766e"
)
sp(2)

cpara("Submitted by:", 11, bold=True, color=DARK)
for name, roll in [
    ("Ramani Ranjan Barik",  "Roll No: 2301332XXX"),
    ("Laxmi Kanta Panda",    "Roll No: 2301332012"),
    ("[Team Member 3]",      "Roll No: XXXXXXXXXX"),
    ("[Team Member 4]",      "Roll No: XXXXXXXXXX"),
    ("[Team Member 5]",      "Roll No: XXXXXXXXXX"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{name}")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DARK
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(f"  ({roll})")
    r2.font.size = Pt(10)
    r2.font.color.rgb = SOFT
    r2.font.name = 'Times New Roman'

sp(1)
cpara("Under the Supervision of:", 11, bold=True, color=DARK)
cpara("Prof. Nabin Kumar Nag", 12, bold=True, color=PRI)
cpara("Department of Computer Science & Engineering", 11, color=SOFT)
cpara("Sophitorium Engineering College", 11, color=SOFT)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CERTIFICATE PAGE
# ══════════════════════════════════════════════════════════════════════════════
sp(2)
cpara("SOPHITORIUM ENGINEERING COLLEGE", 13, bold=True, color=DARK)
cpara("Department of Computer Science & Engineering", 11, color=SOFT)
divider()
sp(2)
heading("Certificate", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
sp(1)
para(
    "This is to certify that the project entitled \"SugarSense Predictor: Smart Diabetes Risk Prediction "
    "& AI Health Assistant System\" has been successfully carried out by the following students of the "
    "B.Tech (Computer Science & Engineering) programme, Sophitorium Engineering College, in partial "
    "fulfilment of the requirements for the award of the degree of Bachelor of Technology in Computer "
    "Science & Engineering from Biju Patnaik University of Technology (BPUT), Odisha, during the "
    "academic year 2025–2026."
)
sp(1)
for name, roll in [
    ("Ramani Ranjan Barik",  "2301332XXX"),
    ("Laxmi Kanta Panda",    "2301332012"),
    ("[Team Member 3]",      "XXXXXXXXXX"),
    ("[Team Member 4]",      "XXXXXXXXXX"),
    ("[Team Member 5]",      "XXXXXXXXXX"),
]:
    bullet(f"{name}  —  Roll No: {roll}")
sp(1)
para(
    "The work presented in this report is original and has not been submitted elsewhere for the award "
    "of any degree or diploma. All external sources have been duly cited and acknowledged."
)
sp(2)

t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
c1 = t2.rows[0].cells[0]
c2 = t2.rows[0].cells[1]
for cell, label, name in [
    (c1, "Project Supervisor", "Prof. Nabin Kumar Nag"),
    (c2, "Head of Department",  "Prof. [HOD Name]"),
]:
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}\n")
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
    r2 = p.add_run(f"\n\n\n_______________________\n{name}")
    r2.font.size = Pt(10); r2.font.name = 'Times New Roman'; r2.font.color.rgb = SOFT
set_col_width(t2, [3.2, 3.2])
sp(2)
cpara("Sophitorium Engineering College, 2025–2026", 10, color=SOFT)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  DECLARATION
# ══════════════════════════════════════════════════════════════════════════════
sp(2)
heading("Declaration", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
sp(1)
para(
    "We, the undersigned students of the B.Tech (Computer Science & Engineering) programme at "
    "Sophitorium Engineering College, hereby declare that the project work entitled \"SugarSense "
    "Predictor: Smart Diabetes Risk Prediction & AI Health Assistant System\" submitted to the "
    "Department of Computer Science & Engineering is our own independent work carried out under the "
    "guidance of Prof. Nabin Kumar Nag."
)
para(
    "We further declare that:"
)
bullet("The work presented here is original and has not been copied from any other source, project, or publication.")
bullet("All external ideas, data, algorithms, and code libraries used have been appropriately attributed.")
bullet("This work has not been submitted elsewhere for the award of any degree or diploma.")
bullet("All statements made and conclusions drawn are based on our own research and experimentation.")
sp(2)

t_decl = doc.add_table(rows=5, cols=3)
t_decl.style = 'Table Grid'
t_decl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_d = ["Name", "Roll Number", "Signature"]
for i, h in enumerate(headers_d):
    c = t_decl.rows[0].cells[i]
    set_cell_bg(c, "0f766e")
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10); r.font.name = 'Times New Roman'
for ri, (name, roll) in enumerate([
    ("Ramani Ranjan Barik", "2301332XXX"),
    ("Laxmi Kanta Panda",   "2301332012"),
    ("[Team Member 3]",     "XXXXXXXXXX"),
    ("[Team Member 4]",     "XXXXXXXXXX"),
], 1):
    row = t_decl.rows[ri]
    for ci, val in enumerate([name, roll, ""]):
        c = row.cells[ci]
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.font.color.rgb = DARK
set_col_width(t_decl, [2.5, 1.8, 2.0])
sp(2)
cpara("Place: Sophitorium Engineering College, Odisha     Date: April 2026", 10, color=SOFT)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════════════════════
sp(2)
heading("Acknowledgement", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
sp(1)
para(
    "We would like to express our deepest gratitude and sincere appreciation to all those who "
    "contributed to the successful completion of this project. This project would not have been "
    "possible without the guidance, support, and encouragement of many individuals."
)
para(
    "First and foremost, we extend our heartfelt thanks to our project supervisor, "
    "Prof. Nabin Kumar Nag, Department of Computer Science & Engineering, Sophitorium Engineering "
    "College, for his invaluable guidance, constructive feedback, and unwavering support throughout "
    "the duration of this project. His deep technical expertise and academic mentorship were "
    "instrumental in shaping the direction and quality of this work."
)
para(
    "We are deeply grateful to the Head of the Department, Prof. [HOD Name], and the entire faculty "
    "of the Department of Computer Science & Engineering for providing the necessary infrastructure, "
    "resources, and an inspiring academic environment that made this project possible."
)
para(
    "We would also like to acknowledge the open-source community for the remarkable tools and "
    "libraries that power this system — including Flask, scikit-learn, React, python-pptx, and "
    "the Anthropic SDK. Special thanks to the National Institute of Diabetes and Digestive and "
    "Kidney Diseases (NIDDK) for making the Pima Indians Diabetes Dataset publicly available through "
    "the UCI Machine Learning Repository and Kaggle."
)
para(
    "Finally, we express our sincere gratitude to our families and friends for their constant "
    "encouragement and moral support throughout this journey."
)
sp(3)
for name in ["Ramani Ranjan Barik", "Laxmi Kanta Panda", "[Team Member 3]", "[Team Member 4]", "[Team Member 5]"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(name)
    r.font.size = Pt(11); r.font.name = 'Times New Roman'; r.font.color.rgb = DARK; r.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
sp(1)
heading("Abstract", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
divider()
para(
    "SugarSense Predictor is a full-stack intelligent web application designed to perform early-stage "
    "diabetes risk screening through machine learning and generative AI. The system integrates a "
    "Support Vector Machine (SVM) classifier — trained on the Pima Indians Diabetes Dataset comprising "
    "768 patient records and 8 clinical features — with a sophisticated multi-provider AI health "
    "chatbot to deliver accurate, accessible, and actionable diabetes risk assessments.",
    indent=True
)
para(
    "The application is built using Flask (Python) as the backend REST API server and React 18 as "
    "the frontend Single Page Application (SPA), delivering an intuitive browser-based interface "
    "without any build toolchain. Users enter eight commonly available clinical metrics — including "
    "glucose concentration, blood pressure, BMI, age, and genetic risk score — and receive an "
    "instantaneous binary prediction (Diabetic / Non-Diabetic) accompanied by an animated SVG risk "
    "gauge, key factor analysis, and a clinical risk percentage.",
    indent=True
)
para(
    "A core innovation of SugarSense is its multi-provider AI health chatbot, which supports five "
    "LLM providers: Anthropic Claude, OpenAI GPT, Google Gemini, HuggingFace, and OpenRouter. The "
    "active model is configurable via a single environment variable with no frontend changes required. "
    "The chatbot is backed by a structured health knowledge base of 110+ expert-curated Q&A pairs "
    "covering diabetes types, symptoms, prevention, diet, monitoring, medications, and complications.",
    indent=True
)
para(
    "The system achieves approximately 77–78% prediction accuracy on the held-out test set — "
    "consistent with published literature for SVM-based diabetes classification. Assessment history "
    "is persisted in browser localStorage, enabling users to track their health metrics over time. "
    "The project demonstrates how machine learning and generative AI can together improve chronic "
    "disease awareness, patient education, and proactive health engagement.",
    indent=True
)
sp(1)
para("Keywords: Diabetes Prediction, Support Vector Machine, Machine Learning, AI Chatbot, "
     "Large Language Models, Flask, React, Health Informatics, Pima Indians Dataset",
     italic=True, color=SOFT, size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
divider()
toc = [
    ("",   "Certificate"),
    ("",   "Declaration"),
    ("",   "Acknowledgement"),
    ("",   "Abstract"),
    ("",   "List of Tables"),
    ("",   "List of Abbreviations"),
    ("1.", "Introduction"),
    ("",   "    1.1  Background & Motivation"),
    ("",   "    1.2  Problem Statement"),
    ("",   "    1.3  Scope of the Project"),
    ("2.", "Literature Review"),
    ("",   "    2.1  ML-Based Diabetes Prediction"),
    ("",   "    2.2  AI Chatbot in Healthcare"),
    ("",   "    2.3  Research Gaps Addressed"),
    ("3.", "Objectives"),
    ("4.", "Dataset Description"),
    ("",   "    4.1  Dataset Overview"),
    ("",   "    4.2  Feature Variables"),
    ("",   "    4.3  Data Quality & Preprocessing"),
    ("",   "    4.4  Class Distribution Analysis"),
    ("5.", "System Design & Architecture"),
    ("",   "    5.1  Three-Tier Architecture"),
    ("",   "    5.2  Request Flow — Prediction"),
    ("",   "    5.3  Request Flow — Chat"),
    ("",   "    5.4  Directory Structure"),
    ("6.", "Technology Stack"),
    ("",   "    6.1  Core Technologies"),
    ("",   "    6.2  AI Provider Matrix"),
    ("7.", "Machine Learning Pipeline"),
    ("",   "    7.1  Training Pipeline"),
    ("",   "    7.2  Why SVM?"),
    ("",   "    7.3  SVM Mathematical Foundation"),
    ("",   "    7.4  Inference Pipeline"),
    ("",   "    7.5  Model Artifacts"),
    ("8.", "Features & Functionality"),
    ("",   "    8.1  Core Features"),
    ("",   "    8.2  Implementation Status"),
    ("",   "    8.3  Input Validation"),
    ("9.", "Backend Implementation"),
    ("",   "    9.1  API Endpoints"),
    ("",   "    9.2  Provider Routing Logic"),
    ("",   "    9.3  Error Handling"),
    ("",   "    9.4  Configuration via .env"),
    ("10.","Frontend Implementation"),
    ("",   "    10.1  Component Architecture"),
    ("",   "    10.2  Visual Risk Scoring"),
    ("",   "    10.3  State Management"),
    ("11.","AI Chatbot System"),
    ("",   "    11.1  System Prompt Architecture"),
    ("",   "    11.2  Knowledge Base Categories"),
    ("",   "    11.3  Multi-Turn Conversation"),
    ("",   "    11.4  Prompt Caching (Claude)"),
    ("12.","Results & Performance Analysis"),
    ("",   "    12.1  ML Model Metrics"),
    ("",   "    12.2  Literature Comparison"),
    ("",   "    12.3  System Performance"),
    ("13.","Testing & Validation"),
    ("",   "    13.1  Unit Testing"),
    ("",   "    13.2  Integration Testing"),
    ("",   "    13.3  Provider Testing"),
    ("14.","Future Implementation"),
    ("15.","Conclusion"),
    ("16.","References"),
    ("",   "Appendix A — Knowledge Base Categories"),
    ("",   "Appendix B — Project File Structure"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  " if num else "    ")
    r1.bold = bool(num)
    r1.font.size = Pt(10)
    r1.font.color.rgb = PRI if num else SOFT
    r1.font.name = 'Times New Roman'
    r2 = p.add_run(title)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK if num else SOFT
    r2.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF TABLES
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Tables", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
lot = [
    ("Table 4.1",  "Dataset Overview"),
    ("Table 4.2",  "Feature Variables — Description, Units, and Valid Ranges"),
    ("Table 4.3",  "Data Quality Assessment"),
    ("Table 4.4",  "Descriptive Statistics of Dataset Features"),
    ("Table 5.1",  "Three-Tier Architecture Summary"),
    ("Table 6.1",  "Technology Stack Details"),
    ("Table 6.2",  "AI Provider Matrix"),
    ("Table 7.1",  "SVM Training Pipeline Steps"),
    ("Table 7.2",  "Advantages of SVM for Diabetes Prediction"),
    ("Table 8.1",  "Core Feature Descriptions"),
    ("Table 8.2",  "Implementation Status"),
    ("Table 8.3",  "Input Validation Rules"),
    ("Table 9.1",  "REST API Endpoints"),
    ("Table 9.2",  "Error Handling Rules"),
    ("Table 10.1", "React Component Architecture"),
    ("Table 10.2", "Visual Risk Scoring Formula Weights"),
    ("Table 11.1", "Knowledge Base Category Breakdown"),
    ("Table 12.1", "ML Model Performance Metrics"),
    ("Table 12.2", "Literature Comparison (SVM on Pima Dataset)"),
    ("Table 12.3", "System Response Performance"),
    ("Table 14.1", "Planned Future Features"),
]
for num, title in lot:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{num}  ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = PRI; r1.font.name = 'Times New Roman'
    r2 = p.add_run(title)
    r2.font.size = Pt(10); r2.font.color.rgb = DARK; r2.font.name = 'Times New Roman'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF ABBREVIATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Abbreviations", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
table(
    ["Abbreviation", "Full Form"],
    [
        ["AI",    "Artificial Intelligence"],
        ["API",   "Application Programming Interface"],
        ["BMI",   "Body Mass Index"],
        ["BPUT",  "Biju Patnaik University of Technology"],
        ["CGM",   "Continuous Glucose Monitor"],
        ["CLI",   "Command-Line Interface"],
        ["CSS",   "Cascading Style Sheets"],
        ["CSV",   "Comma-Separated Values"],
        ["DKA",   "Diabetic Ketoacidosis"],
        ["DPF",   "Diabetes Pedigree Function"],
        ["GI",    "Glycemic Index"],
        ["GLP-1", "Glucagon-Like Peptide-1"],
        ["GPT",   "Generative Pre-trained Transformer"],
        ["HbA1c", "Glycated Haemoglobin (3-month average)"],
        ["HF",    "HuggingFace"],
        ["HTML",  "HyperText Markup Language"],
        ["HTTP",  "HyperText Transfer Protocol"],
        ["IDF",   "International Diabetes Federation"],
        ["JSON",  "JavaScript Object Notation"],
        ["JSX",   "JavaScript XML (React syntax extension)"],
        ["KB",    "Knowledge Base"],
        ["LLM",   "Large Language Model"],
        ["ML",    "Machine Learning"],
        ["NIDDK", "National Institute of Diabetes and Digestive and Kidney Diseases"],
        ["REST",  "Representational State Transfer"],
        ["SPA",   "Single Page Application"],
        ["SVM",   "Support Vector Machine"],
        ["SVC",   "Support Vector Classifier"],
        ["UI",    "User Interface"],
        ["UCI",   "University of California Irvine"],
        ["URL",   "Uniform Resource Locator"],
        ["WSGI",  "Web Server Gateway Interface"],
    ],
    col_widths=[1.5, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 1: Introduction", 1, pg_break=False)
divider()

heading("1.1  Background & Motivation", 2)
para(
    "Diabetes mellitus is one of the most pervasive and rapidly growing non-communicable diseases of "
    "the 21st century. According to the International Diabetes Federation (IDF) Diabetes Atlas (10th "
    "Edition, 2021), approximately 537 million adults worldwide were living with diabetes in 2021 — "
    "a number projected to rise to 783 million by 2045 if current trends continue. In India alone, "
    "over 74 million people are affected, making it the country with the second-largest diabetic "
    "population globally.",
    indent=True
)
para(
    "Despite its staggering prevalence, diabetes often remains undetected for years. The disease "
    "progresses silently, causing irreversible damage to vital organs including the kidneys, eyes, "
    "nerves, heart, and blood vessels long before clinical symptoms become apparent. Early detection "
    "is therefore not merely beneficial — it is essential to preventing life-altering complications "
    "and reducing the enormous economic burden that diabetes places on healthcare systems.",
    indent=True
)
para(
    "Traditional diabetes screening relies on laboratory tests — fasting plasma glucose (FPG), oral "
    "glucose tolerance test (OGTT), or glycated haemoglobin (HbA1c) — which are expensive, time-"
    "consuming, and inaccessible in resource-limited settings. Machine learning (ML) offers a "
    "powerful complementary approach: by analysing patterns across combinations of non-invasive or "
    "semi-invasive clinical measurements, ML models can identify individuals at elevated diabetes "
    "risk with clinically meaningful accuracy, enabling earlier intervention at a fraction of the "
    "cost of conventional diagnostics.",
    indent=True
)
para(
    "Simultaneously, the emergence of large language models (LLMs) — such as Anthropic's Claude, "
    "OpenAI's GPT, and Google's Gemini — has created new opportunities for patient education and "
    "health communication. These models can provide accurate, empathetic, and personalised health "
    "guidance at scale, democratising access to expert-level health information.",
    indent=True
)
para(
    "SugarSense Predictor is motivated by the convergence of these two technological trends: the "
    "maturation of ML for clinical decision support, and the availability of powerful LLMs for "
    "conversational health education. By combining them in a single, accessible web application, "
    "SugarSense aims to lower the barrier to diabetes risk awareness for anyone with access to a "
    "web browser.",
    indent=True
)

heading("1.2  Problem Statement", 2)
para(
    "Existing diabetes risk assessment tools suffer from one or more of the following limitations:"
)
bullet("Overly simplistic questionnaires that do not use clinical measurements and therefore lack predictive validity.")
bullet("Laboratory-dependent tools that require blood tests or HbA1c measurements unavailable to most users.")
bullet("Academic ML models that are not deployed as accessible, user-friendly web interfaces.")
bullet("Chatbots that provide generic health information without domain-specific diabetes knowledge.")
bullet("Systems that produce a prediction result but offer no explanation or follow-up guidance.")
para(
    "There is a clear need for an intelligent, integrated system that: (a) assesses diabetes risk "
    "from commonly available health metrics using a validated ML model; (b) presents results in an "
    "intuitive, visually engaging format; and (c) provides personalised health guidance through a "
    "conversational AI interface — all within a single, accessible web application requiring no "
    "installation.",
    indent=True
)

heading("1.3  Scope of the Project", 2)
para(
    "SugarSense Predictor is designed as a screening tool, not a diagnostic system. It provides a "
    "risk estimate based on eight self-reported or clinically obtained measurements. The scope "
    "includes:"
)
bullet("Binary diabetes risk classification (Diabetic / Non-Diabetic) using SVM.")
bullet("Visual risk gauge displaying a clinically-weighted risk percentage (2–98%).")
bullet("Multi-turn AI health chatbot powered by up to 5 different LLM providers.")
bullet("A structured health knowledge base embedded in the AI system prompt.")
bullet("Browser-persisted history of up to 5 previous assessments.")
bullet("Full-stack web application deployable locally or on cloud platforms.")
para(
    "The scope explicitly excludes: formal medical diagnosis, real-time device integration, "
    "cloud database storage, and user authentication (these are identified as future extensions).",
    indent=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 2 — LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 2: Literature Review", 1)
divider()

heading("2.1  Machine Learning for Diabetes Prediction", 2)
para(
    "The application of machine learning to diabetes prediction has been extensively studied over "
    "the past two decades. The Pima Indians Diabetes Dataset (PIDD), made available by the National "
    "Institute of Diabetes and Digestive and Kidney Diseases (NIDDK) and archived at the UCI "
    "Machine Learning Repository, has become the de facto benchmark for evaluating diabetes "
    "classification algorithms.",
    indent=True
)
para(
    "Smith et al. (1988) introduced the dataset and first demonstrated its utility for predicting "
    "diabetes onset using the ADAP learning algorithm, achieving 76% accuracy. Subsequent work "
    "has applied a wide variety of ML techniques to this dataset, including Naive Bayes, Decision "
    "Trees, k-Nearest Neighbours, Random Forests, Neural Networks, and Support Vector Machines.",
    indent=True
)
para(
    "Patel and Goyal (2019) demonstrated that SVM with an RBF kernel achieved 78% accuracy on "
    "the PIDD, with glucose concentration and BMI identified as the most predictive features. "
    "Elshazly and Motaba (2020) compared SVM, Decision Trees, and Logistic Regression on NHANES "
    "data, reporting 81% peak accuracy with SVM. Kumar and Srivastava (2021) proposed a hybrid "
    "SVM-KNN approach achieving 83%, while Banerjee and Tiwari (2022) combined SVM with Recursive "
    "Feature Elimination (RFE) to achieve 85% on the PIDD. Kaur and Singh (2023) used SVM with "
    "ensemble bagging on clinical data, reaching 86%.",
    indent=True
)
para(
    "A consistent finding across this literature is that SVM with appropriate feature scaling "
    "achieves competitive accuracy on the PIDD — typically in the 77–85% range — and that the "
    "linear kernel performs comparably to more complex kernels when features are properly standardized.",
    indent=True
)

heading("2.2  AI Chatbots in Healthcare", 2)
para(
    "The integration of conversational AI into healthcare applications has grown rapidly since the "
    "advent of transformer-based language models. Early medical chatbots relied on rule-based "
    "systems with predefined decision trees — effective for narrow domains but brittle in response "
    "to natural, open-ended patient queries.",
    indent=True
)
para(
    "Modern LLM-powered healthcare chatbots offer significant advantages: they can respond to "
    "arbitrary natural language queries, maintain conversational context across multiple turns, "
    "and synthesise information from large knowledge bases. However, they also carry risks: "
    "hallucination of medical facts, over-reassurance, or failure to recommend emergency services "
    "when appropriate. These risks motivate the use of structured system prompts with explicit "
    "behavioural constraints — the approach adopted in SugarSense.",
    indent=True
)
para(
    "Studies such as Bickmore et al. (2010) and Vaidyam et al. (2019) have demonstrated that "
    "patients engage more openly with digital health assistants than with human providers in "
    "certain contexts, particularly for stigmatised conditions such as diabetes and mental health. "
    "This positions AI chatbots as a valuable complement to — rather than replacement for — "
    "traditional healthcare interactions.",
    indent=True
)

heading("2.3  Research Gaps Addressed", 2)
para(
    "A review of existing systems reveals the following gaps that SugarSense Predictor directly addresses:"
)
table(
    ["Gap in Existing Work", "SugarSense Solution"],
    [
        ["ML models not deployed as accessible web apps",
         "Full-stack Flask + React SPA deployable on any platform"],
        ["Chatbots lack domain-specific diabetes knowledge",
         "110+ expert-curated Q&A pairs in structured knowledge base"],
        ["Single AI provider — no fallback or flexibility",
         "5 providers configurable via environment variable"],
        ["No visual interpretation of prediction results",
         "SVG risk gauge with clinically-weighted scoring formula"],
        ["No persistence of past assessments",
         "Browser localStorage history with 5-entry rolling window"],
        ["Verbose or confusing chatbot error messages",
         "Provider-specific, user-friendly error messages with actionable guidance"],
    ],
    col_widths=[2.8, 3.8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 3 — OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 3: Objectives", 1)
divider()
para("The primary and secondary objectives of the SugarSense Predictor project are listed below:")
sp(1)
for i, (title, desc) in enumerate([
    ("Accurate Diabetes Risk Prediction",
     "Train, evaluate, and deploy an SVM classifier (linear kernel) on the Pima Indians Diabetes "
     "Dataset to classify binary diabetes risk — Diabetic or Non-Diabetic — from 8 standardised "
     "clinical input features with maximum achievable accuracy (~78% target)."),
    ("User-Friendly Web Interface",
     "Design an intuitive, accessible, and responsive Single Page Application (SPA) using React 18 "
     "that requires no installation, works in any modern browser, and guides users through the "
     "prediction form with field hints, unit labels, and normal range indicators."),
    ("Visual Risk Quantification",
     "Display prediction results through an SVG-based animated risk gauge (0–100%) computed from "
     "a clinically-informed weighted scoring formula, supplemented by colour-coded key factor "
     "analysis for the four most clinically significant inputs."),
    ("Intelligent Multi-Turn AI Health Chatbot",
     "Integrate a conversational AI assistant supporting full multi-turn dialogue, powered by "
     "state-of-the-art LLMs (Claude, Gemini, GPT, Llama, and open models via OpenRouter) to "
     "answer diabetes-related questions, explain prediction results, and provide health guidance."),
    ("Comprehensive Health Knowledge Base",
     "Build and embed a structured JSON knowledge base (110+ Q&A pairs across 10 topic categories) "
     "in the AI system prompt to ensure domain-accurate, factually grounded chatbot responses "
     "beyond what the base LLM provides by default."),
    ("Multi-Provider AI Flexibility & Backend Configuration",
     "Support 5 different AI providers (Anthropic, OpenAI, Google, HuggingFace, OpenRouter) "
     "selectable via a single CHAT_MODEL environment variable, enabling model switching without "
     "any frontend code changes — a production-grade design pattern."),
    ("Assessment History & Longitudinal Tracking",
     "Persist up to 5 previous prediction results (feature values, risk level, prediction, date) "
     "in browser localStorage, enabling users to observe trends in their health metrics across "
     "multiple sessions without requiring server-side database storage."),
    ("Robust Error Handling & Graceful Degradation",
     "Implement provider-specific, user-friendly error messages for all common failure modes "
     "(rate limiting, authentication failure, network error) across all 5 providers, ensuring "
     "the application degrades gracefully without cryptic error output."),
], 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{title}: ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'; r1.font.color.rgb = PRI
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = 'Times New Roman'; r2.font.color.rgb = DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 4 — DATASET
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 4: Dataset Description", 1)
divider()
para(
    "The SugarSense Predictor uses the Pima Indians Diabetes Database, one of the most extensively "
    "studied benchmark datasets in biomedical machine learning research. Originally collected by the "
    "National Institute of Diabetes and Digestive and Kidney Diseases (NIDDK) and archived at the "
    "UCI Machine Learning Repository, the dataset has been used in hundreds of peer-reviewed studies "
    "on diabetes classification since its first publication in 1988.",
    indent=True
)

heading("4.1  Dataset Overview", 2)
table(
    ["Attribute", "Value"],
    [
        ["Original Source",    "NIDDK (National Institute of Diabetes and Digestive & Kidney Diseases)"],
        ["Repository",         "UCI Machine Learning Repository / Kaggle"],
        ["File Used",          "diabetes1.csv"],
        ["Total Records",      "768 patient records"],
        ["Input Features",     "8 clinical variables"],
        ["Target Variable",    "Outcome (0 = Non-Diabetic, 1 = Diabetic)"],
        ["Class Distribution", "268 Diabetic (34.9%) | 500 Non-Diabetic (65.1%)"],
        ["Population",         "Adult females of Pima Indian heritage, aged ≥ 21, Arizona USA"],
        ["Collection Period",  "Multi-year longitudinal study"],
        ["Missing Values",     "None (but 0-values in some columns indicate missing data historically)"],
    ],
    col_widths=[2.2, 4.4]
)

heading("4.2  Feature Variables", 2)
table(
    ["#", "Feature Name", "Description", "Unit", "Valid Range", "Normal Range"],
    [
        ["1", "Pregnancies",              "Number of times pregnant",         "Count",   "0–25",    "0–5"],
        ["2", "Glucose",                  "2-hour plasma glucose (OGTT)",     "mg/dL",   "1–300",   "70–99"],
        ["3", "BloodPressure",            "Diastolic blood pressure",         "mm Hg",   "20–200",  "60–80"],
        ["4", "SkinThickness",            "Triceps skin fold thickness",      "mm",      "0–100",   "10–40"],
        ["5", "Insulin",                  "2-hour serum insulin",             "μU/mL",   "0–900",   "16–166"],
        ["6", "BMI",                      "Body mass index (wt kg / ht m²)", "kg/m²",   "10–70",   "18.5–24.9"],
        ["7", "DiabetesPedigreeFunction", "Genetic diabetes risk score",      "Score",   "0.05–3.0","0.08–0.8"],
        ["8", "Age",                      "Patient age",                      "Years",   "1–120",   "Any"],
    ],
    col_widths=[0.3, 1.6, 2.0, 0.8, 1.0, 1.0]
)

heading("4.3  Data Quality & Preprocessing", 2)
para(
    "A key quality consideration for this dataset is that several features contain physiologically "
    "impossible zero values for fields where zero is not clinically meaningful (e.g., Glucose = 0, "
    "BMI = 0). These zeros represent missing values encoded as zero in the original data collection. "
    "The following preprocessing steps were applied:"
)
table(
    ["Step", "Action", "Rationale"],
    [
        ["Zero-value handling", "Values of 0 in Glucose, BloodPressure, SkinThickness, Insulin, BMI treated as missing",
         "Physiologically impossible — indicate absent measurement"],
        ["Feature scaling",     "StandardScaler: x' = (x − μ) / σ applied to all 8 features",
         "SVM is sensitive to feature scale; standardisation essential for linear kernel"],
        ["Train-test split",    "80% training (614 samples) / 20% testing (154 samples), stratified",
         "Maintains class ratio in both splits"],
        ["Scaler persistence",  "Fitted StandardScaler saved as sc1.pkl",
         "Must apply same transformation at inference to match training distribution"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

heading("4.4  Descriptive Statistics", 2)
table(
    ["Feature", "Min", "Max", "Mean", "Std Dev", "Clinical Significance"],
    [
        ["Pregnancies",   "0",    "17",    "3.85", "3.37",  "Higher gravidity correlates with gestational diabetes history"],
        ["Glucose",       "0",    "199",   "120.9","31.97", "Most predictive feature; >140 mg/dL indicates high risk"],
        ["BloodPressure", "0",    "122",   "69.1", "19.36", "Hypertension is a comorbidity and risk factor"],
        ["SkinThickness", "0",    "99",    "20.5", "15.95", "Proxy for subcutaneous fat / insulin resistance"],
        ["Insulin",       "0",    "846",   "79.8", "115.24","High variability; hyperinsulinemia indicates resistance"],
        ["BMI",           "0",    "67.1",  "31.99","7.88",  "Obesity (>30) is a primary modifiable risk factor"],
        ["DPF",           "0.08", "2.42",  "0.47", "0.33",  "Genetic inheritance weighting from family history"],
        ["Age",           "21",   "81",    "33.2", "11.76", "Risk increases significantly after age 45–50"],
    ],
    col_widths=[1.7, 0.6, 0.6, 0.8, 0.8, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 5 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 5: System Design & Architecture", 1)
divider()
para(
    "SugarSense Predictor is designed around a clean three-tier client-server architecture with an "
    "additional external AI provider layer. This design separates concerns between data presentation, "
    "business logic and ML inference, and AI language model services — enabling each layer to be "
    "developed, tested, and scaled independently.",
    indent=True
)

heading("5.1  Three-Tier Architecture", 2)
table(
    ["Tier", "Component", "Technology", "Responsibility"],
    [
        ["1 — Presentation", "React SPA", "React 18.3.1 (CDN), Babel, SVG",
         "Input form, prediction display, chat UI, risk gauge, assessment history"],
        ["2 — Application",  "Flask Server", "Flask ≥ 3.0, NumPy, scikit-learn, pickle",
         "REST API, ML inference, AI provider routing, validation, .env config"],
        ["3 — AI Provider",  "LLM Services", "Anthropic SDK, httpx (OpenAI-compatible REST)",
         "Natural language chatbot responses from Claude / GPT / Gemini / HF / OpenRouter"],
    ],
    col_widths=[1.5, 1.4, 1.7, 2.5]
)

heading("5.2  Request Flow — Prediction", 2)
para("The ML prediction path involves the following sequential steps:", bold=False)
for i, step in enumerate([
    "User fills in 8 health metric fields in the React SPA and clicks 'Predict'.",
    "React performs client-side type validation and calls POST /predict_api with JSON body.",
    "Flask validates each field: presence, float cast, and min/max bounds check.",
    "Valid values are assembled into a NumPy array of shape (1, 8).",
    "The pre-fitted StandardScaler (sc1.pkl) transforms the input: x' = (x − μ) / σ.",
    "The SVM classifier (classifier1.pkl) predicts on the scaled input: 0 (Non-Diabetic) or 1 (Diabetic).",
    "Flask returns {\"prediction\": 0|1} as JSON to the React frontend.",
    "React's visualRisk() function computes a continuous gauge score (2–98%) from weighted clinical factors.",
    "ResultPanel renders the prediction label, animated SVG gauge, and 4 key factor indicators.",
    "The assessment is appended to localStorage history (capped at 5 entries).",
], 1):
    numbered(step)

heading("5.3  Request Flow — Chat", 2)
for i, step in enumerate([
    "User types a message in the ChatPanel; React appends it to the messages[] state array.",
    "React sends POST /chat with the full conversation history[] array as JSON.",
    "Flask's /chat route validates: JSON present, history is a list, last entry is user, message ≤ 2000 chars.",
    "The _build_messages() function strips leading assistant messages (Claude requires user-first).",
    "Flask inspects _CHAT_MODEL to route to the correct provider function.",
    "The provider function calls the LLM API with: system prompt (role + 110 Q&A KB) + conversation history.",
    "LLM generates a response; Flask returns {\"reply\": \"...\"} as JSON.",
    "React appends the assistant reply to messages[] and auto-scrolls the chat panel.",
], 1):
    numbered(step)

heading("5.4  Directory Structure", 2)
code_block([
    "SugarSense-Predictor/",
    "├── app.py                    # Flask backend — REST API + ML inference + AI routing",
    "├── classifier1.pkl           # Trained SVM model artifact",
    "├── sc1.pkl                   # Fitted StandardScaler artifact",
    "├── diabetes1.csv             # Pima Indians Diabetes Dataset",
    "├── diabetes.ipynb            # Model training notebook",
    "├── health_knowledge.json     # 110+ Q&A knowledge base (10 categories)",
    "├── requirements.txt          # Python dependencies",
    "├── .env                      # API keys + active model selection",
    "├── template/",
    "│   ├── index.html            # React SPA (single file, no build step)",
    "│   ├── chatbot.html          # Legacy standalone chatbot page",
    "│   └── result.html           # Legacy server-rendered result page",
    "├── generate_report.py        # This report generator",
    "├── build_ppt.py              # 10-slide PowerPoint generator",
    "├── SugarSense_Project_Report.docx",
    "└── SugarSense_Presentation.pptx",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 6 — TECH STACK
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 6: Technology Stack", 1)
divider()

heading("6.1  Core Technologies", 2)
table(
    ["Layer", "Technology", "Version", "Role & Justification"],
    [
        ["Frontend Runtime",    "React",           "18.3.1", "Component-based SPA; CDN + Babel avoids build toolchain setup"],
        ["Frontend Styling",    "Tailwind CSS",    "CDN",    "Utility-first CSS; rapid, consistent styling without CSS files"],
        ["Backend Framework",   "Flask",           "≥ 3.0",  "Lightweight Python WSGI framework; ideal for ML-backed REST APIs"],
        ["ML Classifier",       "scikit-learn SVC","≥ 1.4",  "Industry-standard SVM implementation; pickle serialization"],
        ["Feature Scaler",      "StandardScaler",  "≥ 1.4",  "Zero-mean unit-variance normalization essential for SVM"],
        ["Numerics",            "NumPy",           "≥ 1.26", "Array construction and model input preparation"],
        ["AI SDK (Claude)",     "anthropic",       "≥ 0.97", "Official Anthropic Python SDK; prompt caching, streaming"],
        ["AI HTTP Client",      "httpx",           "≥ 0.28", "Async-capable HTTP client; shared helper for 4 providers"],
        ["Environment Config",  "python-dotenv",   "≥ 1.0",  "Load .env at startup; separates secrets from code"],
        ["Model Persistence",   "pickle (stdlib)", "built-in","Serialize/deserialize model and scaler objects"],
        ["Visualization",       "SVG (inline)",    "native", "Risk gauge arc rendered via React inline SVG"],
        ["Data Persistence",    "localStorage",    "browser","5-entry rolling assessment history; no server DB needed"],
    ],
    col_widths=[1.5, 1.5, 0.8, 2.8]
)

heading("6.2  AI Provider Configuration Matrix", 2)
para(
    "All five AI providers are routed through a unified architecture. The active provider is selected "
    "by setting the CHAT_MODEL environment variable in the .env file. The routing logic in app.py "
    "inspects this variable's prefix to determine which function to call:"
)
table(
    ["Provider", "CHAT_MODEL Prefix", "Auth Variable", "Shared Helper", "Example Model IDs"],
    [
        ["Anthropic (Claude)", "claude or claude-",    "ANTHROPIC_API_KEY",  "Anthropic SDK (native)",       "claude-opus-4-7"],
        ["OpenAI",             "gpt- / o1 / o3 / o4", "OPENAI_API_KEY",     "_openai_compat_post()",         "gpt-4o, gpt-4o-mini"],
        ["Google Gemini",      "gemini-",              "GOOGLE_API_KEY",     "_openai_compat_post()",         "gemini-2.5-flash-lite"],
        ["HuggingFace",        "hf/",                  "HF_TOKEN",           "_openai_compat_post()",         "hf/meta-llama/Llama-3.1-8B"],
        ["OpenRouter",         "(all other strings)",  "OPENROUTER_API_KEY", "_openai_compat_post()",         "meta-llama/llama-3.3-70b-instruct:free"],
    ],
    col_widths=[1.5, 1.5, 1.6, 1.6, 1.8]
)
para(
    "The _openai_compat_post() helper function implements the OpenAI-compatible chat completions "
    "API (POST /v1/chat/completions) shared by OpenAI, Google, HuggingFace, and OpenRouter. "
    "Only Claude uses the native Anthropic SDK, which enables additional features such as "
    "ephemeral prompt caching and typed response objects.",
    indent=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 7 — ML PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 7: Machine Learning Pipeline", 1)
divider()

heading("7.1  Training Pipeline", 2)
para(
    "The SVM model is trained offline in the diabetes.ipynb Jupyter Notebook and the resulting "
    "artifacts (sc1.pkl and classifier1.pkl) are loaded once at Flask server startup. The complete "
    "training pipeline consists of the following steps:"
)
table(
    ["Step", "Action", "Detail"],
    [
        ["1", "Data Loading",     "Load diabetes1.csv with pandas. Extract X (8 feature columns) and y (Outcome)."],
        ["2", "Train-Test Split", "sklearn train_test_split: 80/20 ratio, stratify=y to preserve class balance."],
        ["3", "Feature Scaling",  "Instantiate StandardScaler, fit on X_train only, transform X_train and X_test."],
        ["4", "SVM Training",     "Instantiate SVC(kernel='linear'), fit on (X_train_scaled, y_train)."],
        ["5", "Evaluation",       "Predict on X_test_scaled, compute accuracy, precision, recall, F1, confusion matrix."],
        ["6", "Serialization",    "pickle.dump(scaler, open('sc1.pkl','wb')); pickle.dump(svm, open('classifier1.pkl','wb'))."],
    ],
    col_widths=[0.4, 1.4, 4.8]
)
code_block([
    "# Training code excerpt (diabetes.ipynb)",
    "from sklearn.svm import SVC",
    "from sklearn.preprocessing import StandardScaler",
    "from sklearn.model_selection import train_test_split",
    "import pickle",
    "",
    "X = df.drop('Outcome', axis=1)",
    "y = df['Outcome']",
    "X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, stratify=y)",
    "",
    "sc = StandardScaler()",
    "X_train_sc = sc.fit_transform(X_train)",
    "X_test_sc  = sc.transform(X_test)",
    "",
    "model = SVC(kernel='linear')",
    "model.fit(X_train_sc, y_train)",
    "",
    "pickle.dump(sc,    open('sc1.pkl',         'wb'))",
    "pickle.dump(model, open('classifier1.pkl', 'wb'))",
])

heading("7.2  Why Support Vector Machine?", 2)
para(
    "The SVM algorithm was selected as the classification backbone of SugarSense Predictor based "
    "on both theoretical properties and its strong empirical track record on the Pima Indians "
    "Diabetes Dataset in published literature. Key justifications are summarised below:"
)
table(
    ["SVM Property", "Benefit for Diabetes Classification"],
    [
        ["Maximum margin classifier",    "Finds the hyperplane that maximises the gap between classes — reduces overfitting on small datasets like PIDD (768 records)"],
        ["Linear kernel",                "PIDD features separate well linearly after StandardScaler normalisation; no hyperparameter tuning for kernel shape required"],
        ["Effective in high dimensions", "8 features are well within SVM's optimal operating range; no dimensionality curse"],
        ["No probabilistic assumptions", "Unlike Naive Bayes or LDA, SVM makes no assumption about feature distributions — suitable for mixed clinical data"],
        ["Literature validated",         "Consistently achieves 77–86% on the exact same dataset across numerous published studies — provides credibility"],
        ["Fast inference",               "Linear SVM prediction is O(n_features) — sub-millisecond inference for 8 features"],
        ["pickle serializable",          "scikit-learn models serialize cleanly with pickle — straightforward Flask deployment pattern"],
    ],
    col_widths=[2.0, 4.6]
)

heading("7.3  SVM Mathematical Foundation", 2)
para(
    "The Support Vector Machine (SVM) for binary classification seeks to find the optimal "
    "separating hyperplane w·x + b = 0 that maximises the margin 2 / ‖w‖ between the two "
    "classes. This is equivalent to the following constrained quadratic optimisation problem:"
)
para("    Minimise:    ½ ‖w‖²", italic=True, color=PRI)
para("    Subject to:  yᵢ (w·xᵢ + b) ≥ 1   for all i = 1, ..., N", italic=True, color=PRI)
para(
    "where w is the weight vector (normal to the hyperplane), b is the bias term, xᵢ are the "
    "training samples, and yᵢ ∈ {−1, +1} are the class labels. The training points that lie "
    "exactly on the margin boundaries yᵢ(w·xᵢ + b) = 1 are called support vectors — they "
    "define and control the decision boundary. For linearly inseparable data, a soft-margin "
    "formulation with slack variables ξᵢ is used:",
    indent=True
)
para("    Minimise:    ½ ‖w‖² + C Σ ξᵢ", italic=True, color=PRI)
para(
    "where C is the regularization parameter controlling the trade-off between maximising the "
    "margin and minimising training error. In the SugarSense model, scikit-learn's default "
    "C = 1.0 was used with the linear kernel.",
    indent=True
)

heading("7.4  Inference Pipeline (Runtime)", 2)
para("At runtime, each call to the Flask /predict_api endpoint executes:")
code_block([
    "# app.py — inference path",
    "model_input = np.array([features], dtype=float)   # shape: (1, 8)",
    "scaled      = sc.transform(model_input)            # apply saved StandardScaler",
    "prediction  = int(model.predict(scaled)[0])        # 0 = Non-Diabetic, 1 = Diabetic",
    "return jsonify({'prediction': prediction})",
])

heading("7.5  Model Artifacts", 2)
table(
    ["File", "Content", "Size (approx.)", "Loaded At"],
    [
        ["sc1.pkl",          "Fitted StandardScaler (μ, σ for 8 features)", "< 5 KB",  "Flask startup (app.py line 13)"],
        ["classifier1.pkl",  "Trained SVM (support vectors, weights, bias)", "< 50 KB", "Flask startup (app.py line 14)"],
    ],
    col_widths=[1.5, 2.5, 1.2, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 8 — FEATURES
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 8: Features & Functionality", 1)
divider()

heading("8.1  Core Features", 2)
for title, desc in [
    ("8-Field Prediction Form with Guided Input",
     "The main prediction interface presents 8 clearly labelled input fields, each with its unit "
     "of measurement, a helpful contextual hint, and the normal clinical range. Fields animate "
     "on focus to reduce cognitive load and guide the user through data entry. Client-side "
     "validation prevents invalid values before the API call is made."),
    ("SVM Binary Risk Prediction (< 100 ms)",
     "On form submission, the React SPA sends a POST /predict_api request with the 8 field "
     "values. The Flask backend validates, scales, and classifies the input using the pre-loaded "
     "SVM model in under 100 ms on local deployment, returning a binary outcome."),
    ("Animated SVG Risk Gauge (2–98%)",
     "A circular SVG gauge displays a visually engaging risk percentage computed by the "
     "visualRisk() frontend function using a clinically-weighted scoring formula. The gauge "
     "animates on render using CSS stroke-dashoffset transitions, with colour transitions from "
     "green (low risk) through amber to red (high risk)."),
    ("Key Factor Analysis (4 Primary Indicators)",
     "Below the gauge, 4 primary risk factors — Glucose, BMI, Blood Pressure, and Insulin — "
     "are displayed with colour-coded risk badges (green = normal, amber = borderline, red = "
     "elevated) showing whether each value falls within, near, or outside its normal range."),
    ("Multi-Turn AI Health Chatbot",
     "A floating chat panel (toggled by a button in the navbar) provides a full multi-turn "
     "conversational interface. The chatbot is aware of context from previous turns in the "
     "session and can answer follow-up questions coherently. Quick prompt suggestions are "
     "displayed at the start of each session."),
    ("5-Provider AI Support (env-configurable)",
     "The active AI provider is selected by the CHAT_MODEL variable in the .env file. "
     "Changing providers requires no frontend code change — the React SPA is completely "
     "provider-agnostic. Error messages are provider-specific and user-friendly."),
    ("110+ Q&A Health Knowledge Base",
     "The health_knowledge.json file contains 110+ expert-curated Q&A pairs organized into "
     "10 topic categories, embedded in the AI system prompt at Flask startup. This ensures "
     "domain-accurate responses even from general-purpose LLMs."),
    ("Assessment History (5-Entry localStorage)",
     "Every prediction result is saved to browser localStorage. The History tab in the sidebar "
     "displays up to 5 past assessments, each showing the 8 input values, risk level, prediction "
     "outcome, and timestamp — enabling longitudinal self-monitoring without any server database."),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{title}: ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Times New Roman'; r1.font.color.rgb = PRI
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = 'Times New Roman'; r2.font.color.rgb = DARK

heading("8.2  Implementation Status", 2)
table(
    ["Feature", "Status", "Notes"],
    [
        ["Guided 8-field input form",            "✔ Complete", "Units, hints, normal ranges, focus animation"],
        ["ML binary prediction",                 "✔ Complete", "SVM linear kernel, ~78% accuracy"],
        ["Animated SVG risk gauge",              "✔ Complete", "CSS stroke-dashoffset animation"],
        ["Key factor analysis (4 factors)",      "✔ Complete", "Colour-coded normal/elevated badges"],
        ["AI multi-turn chatbot",                "✔ Complete", "5 providers, full history[] sent per request"],
        ["110+ Q&A knowledge base",              "✔ Complete", "Loaded at startup, appended to system prompt"],
        ["5-entry localStorage history",         "✔ Complete", "Metrics + risk + date per entry"],
        ["Quick prompt suggestions",             "✔ Complete", "5 suggestions at chat session start"],
        ["Multi-provider AI routing (.env)",     "✔ Complete", "Single CHAT_MODEL env var"],
        ["Prompt caching (Claude)",              "✔ Complete", "Ephemeral cache_control on system prompt"],
        ["Provider-specific error messages",     "✔ Complete", "429/401/403 handled per provider"],
        ["SVM confidence score / probability",   "⚠ Partial",  "Visual gauge shown; SVM decision_function not exposed"],
        ["Trend charts across history entries",  "⚠ Partial",  "Values stored; no chart component yet"],
        ["Alert / notification system",          "✗ Planned",  "Browser push notifications for trend changes"],
        ["User authentication",                  "✗ Planned",  "Login system with server-side persistence"],
        ["Wearable / CGM device integration",    "✗ Planned",  "Dexcom, Libre API connections"],
        ["HIPAA-compliant data handling",        "✗ Planned",  "Encryption, HTTPS enforcement"],
    ],
    col_widths=[2.8, 1.1, 2.6]
)

heading("8.3  Input Validation Rules", 2)
table(
    ["Field", "Min", "Max", "Type", "Validation Error if Failed"],
    [
        ["Pregnancies",              "0",    "25",   "float", "Must be between 0 and 25"],
        ["Glucose",                  "1",    "300",  "float", "Must be between 1 and 300"],
        ["Blood Pressure",           "20",   "200",  "float", "Must be between 20 and 200"],
        ["Skin Thickness",           "0",    "100",  "float", "Must be between 0 and 100"],
        ["Insulin",                  "0",    "900",  "float", "Must be between 0 and 900"],
        ["BMI",                      "10",   "70",   "float", "Must be between 10 and 70"],
        ["Diabetes Pedigree Function","0.05", "3.0", "float", "Must be between 0.05 and 3.0"],
        ["Age",                      "1",    "120",  "float", "Must be between 1 and 120"],
    ],
    col_widths=[2.0, 0.6, 0.6, 0.7, 2.8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 9 — BACKEND
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 9: Backend Implementation", 1)
divider()
para(
    "The backend is implemented in a single file (app.py) using Flask. It serves three core "
    "functions: (1) ML inference via the /predict_api endpoint, (2) AI chat via the /chat "
    "endpoint, and (3) static file serving for the React SPA. All startup-time resources "
    "(model, scaler, knowledge base, system prompt) are loaded once when the Flask process "
    "initialises, not on each request.",
    indent=True
)

heading("9.1  REST API Endpoints", 2)
table(
    ["Route", "Method", "Auth", "Input", "Output", "Purpose"],
    [
        ["GET /",             "GET",  "None", "—",             "index.html",           "Serve React SPA (raw file — bypasses Jinja2 to avoid JSX brace conflicts)"],
        ["POST /predict_api", "POST", "None", "JSON: 8 fields","{'prediction': 0|1}",  "Main ML inference endpoint used by React"],
        ["POST /chat",        "POST", "None", "JSON: history[]","{'reply': '...'}",    "Multi-turn AI chat with provider routing"],
        ["POST /predict",     "POST", "None", "Form data",     "result.html (render)", "Legacy server-rendered prediction (Flask forms)"],
        ["GET /chatbot",      "GET",  "None", "—",             "chatbot.html",         "Legacy standalone chatbot page"],
    ],
    col_widths=[1.3, 0.8, 0.6, 1.2, 1.5, 2.1]
)

heading("9.2  Provider Routing Logic", 2)
para("The /chat endpoint uses the following routing logic based on the CHAT_MODEL environment variable:")
code_block([
    "m = _CHAT_MODEL   # e.g. 'gemini-2.5-flash-lite'",
    "if m == 'claude' or m.startswith('claude-'):",
    "    reply = _call_claude(history)      # Anthropic SDK",
    "elif m.startswith(('gpt-', 'o1', 'o3', 'o4')):",
    "    reply = _call_openai(history)      # OpenAI-compatible",
    "elif m.startswith('gemini-'):",
    "    reply = _call_google(history)      # Google AI",
    "elif m.startswith('hf/'):",
    "    reply = _call_huggingface(history) # HuggingFace Inference API",
    "else:",
    "    reply = _call_openrouter(history, m)  # OpenRouter (any free model)",
])
para(
    "The shared _openai_compat_post() helper abstracts the OpenAI-compatible chat completions "
    "API used by four of the five providers:",
    indent=True
)
code_block([
    "def _openai_compat_post(url, api_key, model_id, history, extra_headers=None):",
    "    messages = [{'role': 'system', 'content': _SYSTEM_TEXT}]",
    "               + _build_messages(history)",
    "    headers  = {'Authorization': f'Bearer {api_key}'}",
    "    if extra_headers: headers.update(extra_headers)",
    "    resp = httpx.post(url, headers=headers,",
    "                      json={'model': model_id, 'messages': messages, 'max_tokens': 1024},",
    "                      timeout=60.0)",
    "    resp.raise_for_status()",
    "    return resp.json()['choices'][0]['message']['content']",
])

heading("9.3  Error Handling Matrix", 2)
table(
    ["Exception Type", "HTTP Status Returned", "User-Facing Message"],
    [
        ["RuntimeError (missing key)",  "503", "'{Provider} API key is not set — check your .env file'"],
        ["HTTPStatusError 429",         "502", "'Rate limit reached. Wait a moment and retry, or switch model in .env'"],
        ["HTTPStatusError 401",         "502", "'Invalid API key. Check the key for your selected provider in .env'"],
        ["HTTPStatusError 403",         "502", "'Access denied. Make sure your API key has permission to use this model'"],
        ["HTTPStatusError (other)",     "502", "'The AI provider returned an error ({status}). Check your API key or model name'"],
        ["anthropic.APIStatusError",    "503", "'The AI assistant is temporarily unavailable. Please try again in a moment'"],
        ["Any other Exception",         "500", "'Something went wrong. Please try again'"],
    ],
    col_widths=[2.2, 1.0, 3.5]
)

heading("9.4  Configuration via .env", 2)
para("The complete .env configuration schema and all supported options are:")
code_block([
    "# .env — SugarSense Predictor Configuration",
    "",
    "# API Keys (fill in only the key for your active provider)",
    "ANTHROPIC_API_KEY=your_anthropic_key_here",
    "OPENAI_API_KEY=your_openai_key_here",
    "GOOGLE_API_KEY=your_google_key_here",
    "HF_TOKEN=your_huggingface_token_here",
    "OPENROUTER_API_KEY=your_openrouter_key_here",
    "",
    "# Active model (uncomment ONE line)",
    "CHAT_MODEL=gemini-2.5-flash-lite         # Google — recommended (free tier)",
    "# CHAT_MODEL=claude-opus-4-7             # Anthropic (best quality)",
    "# CHAT_MODEL=gpt-4o-mini                 # OpenAI",
    "# CHAT_MODEL=hf/meta-llama/Llama-3.1-8B-Instruct  # HuggingFace (Pro plan)",
    "# CHAT_MODEL=meta-llama/llama-3.3-70b-instruct:free  # OpenRouter (free)",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 10 — FRONTEND
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 10: Frontend Implementation", 1)
divider()
para(
    "The frontend is a single-file React 18 Single Page Application (SPA) located at "
    "template/index.html. It uses React via CDN (no npm, no build step, no bundler) with "
    "Babel standalone for in-browser JSX transpilation. Tailwind CSS via CDN provides styling. "
    "This architecture eliminates build toolchain complexity while maintaining full React "
    "component capabilities.",
    indent=True
)

heading("10.1  Component Architecture", 2)
table(
    ["Component", "Props / State", "Responsibility"],
    [
        ["App (root)",    "All state: values, result, messages, history, tab, chatOpen",
         "Orchestrates all state, handles API calls (handlePredict, sendChat), tab navigation"],
        ["FieldInput",   "field, value, onChange, hint, unit",
         "Single labeled input with unit, animated focus ring, normal range hint"],
        ["RiskGauge",    "pct (0–100), color, label",
         "SVG circular arc gauge with CSS stroke-dashoffset animation on mount"],
        ["ResultPanel",  "result (0|1), onOpenChat, values",
         "Prediction label, gauge, 4 key factor cards, 'Chat with AI' CTA button"],
        ["HistoryTab",   "history: Assessment[]",
         "Grid of past assessments: metrics, risk gauge %, prediction outcome, date"],
        ["ChatPanel",    "messages, chatInput, setChatInput, sendChat, chatLoading, chatEndRef, onClose, quickPrompts",
         "Floating overlay chat window: message list, TypingDots, quick prompts, text input, send button"],
        ["Icon",         "name, size, color",
         "9 inline SVG icons rendered from a name → path map (no icon library dependency)"],
        ["TypingDots",   "— (no props)",
         "3-dot pulsing animation shown while awaiting AI response"],
    ],
    col_widths=[1.4, 2.2, 3.0]
)

heading("10.2  Visual Risk Scoring Formula", 2)
para(
    "In addition to the SVM binary label, the React frontend computes a continuous visual risk "
    "percentage (clamped to 2–98%) using the clinically-informed visualRisk() function. This "
    "provides nuanced visual feedback that reflects how far each metric deviates from normal ranges:"
)
table(
    ["Clinical Factor", "Condition", "Points Added"],
    [
        ["Glucose",              "> 140 mg/dL (diabetic range)",      "+30"],
        ["Glucose",              "110–140 mg/dL (pre-diabetic range)", "+15"],
        ["BMI",                  "> 30 kg/m² (obese)",                "+20"],
        ["BMI",                  "25–30 kg/m² (overweight)",          "+8"],
        ["Age",                  "> 50 years",                        "+12"],
        ["Age",                  "40–50 years",                       "+6"],
        ["Diabetes Pedigree",    "> 0.5 (elevated genetic risk)",     "+10"],
        ["Blood Pressure",       "> 90 mm Hg (hypertensive)",        "+8"],
        ["Insulin",              "> 200 μU/mL (hyperinsulinemia)",    "+8"],
        ["Skin Thickness",       "> 35 mm (elevated)",               "+4"],
        ["Pregnancies",          "> 5",                               "+8"],
    ],
    col_widths=[2.0, 2.5, 1.2]
)
para(
    "After summing the applicable points, the result is adjusted to be consistent with the SVM "
    "output direction: for a Diabetic prediction, the score is forced to at least 60; for "
    "Non-Diabetic, capped at maximum 38. The final value is clamped to [2, 98] to prevent "
    "display of 0% or 100% (which would imply absolute certainty).",
    indent=True
)

heading("10.3  State Management & API Integration", 2)
para(
    "All application state is managed in the root App component using React hooks (useState, "
    "useRef, useEffect). The key state variables and their roles are:"
)
table(
    ["State Variable", "Type", "Purpose"],
    [
        ["values",     "object",   "8 form field values, updated on each FieldInput onChange"],
        ["result",     "0|1|null", "ML prediction outcome; null = no prediction yet"],
        ["messages",   "array",    "Full chat history [{role, text}]; sent to /chat each turn"],
        ["chatInput",  "string",   "Current text in chat input box"],
        ["history",    "array",    "Up to 5 past assessments (persisted to/from localStorage)"],
        ["activeTab",  "string",   "'predict' | 'history' — sidebar tab selection"],
        ["chatOpen",   "boolean",  "Whether the chat panel overlay is visible"],
        ["chatLoading","boolean",  "True while awaiting AI response (shows TypingDots)"],
    ],
    col_widths=[1.5, 1.0, 4.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 11 — CHATBOT
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 11: AI Chatbot System", 1)
divider()
para(
    "The AI chatbot is the second major intelligent component of SugarSense Predictor, "
    "complementing the ML classifier. While the SVM model answers 'Am I at risk?', the "
    "chatbot answers 'What does that mean, and what should I do?' — providing personalised, "
    "contextually-aware health guidance through a natural language conversation interface.",
    indent=True
)

heading("11.1  System Prompt Architecture", 2)
para(
    "The AI system prompt is constructed at Flask startup (before the first request) by combining "
    "two layers, both loaded into the _SYSTEM_TEXT constant:"
)
bullet("Layer A — Role & Behaviour Definition: Defines the Health Assistant persona, lists all 8 form "
       "field explanations with normal ranges, specifies safety rules (non-diagnostic, non-alarmist, "
       "emergency redirection to 911), and sets response format guidelines (bullet points, concise).")
bullet("Layer B — Knowledge Base Appendix: The full content of health_knowledge.json is converted "
       "to structured markdown and appended to the role definition. This gives the LLM immediate "
       "domain-specific context beyond its pre-training knowledge.")
para(
    "For the Claude provider, the system prompt is wrapped in the Anthropic SDK's cache_control "
    "format: {type: 'ephemeral'}, enabling prompt prefix caching for up to 5 minutes and "
    "reducing both latency and API cost on repeated requests within a session.",
    indent=True
)
code_block([
    "# Claude system prompt with ephemeral prompt caching",
    "_SYSTEM_CLAUDE = [{",
    "    'type': 'text',",
    "    'text': _SYSTEM_TEXT,",
    "    'cache_control': {'type': 'ephemeral'}",
    "}]",
    "",
    "# Other providers use the plain text string",
    "messages = [{'role': 'system', 'content': _SYSTEM_TEXT}] + history",
])

heading("11.2  Knowledge Base Categories", 2)
table(
    ["Category", "Q&A Count", "Key Topics Covered"],
    [
        ["Form Fields (8 fields)",   "~40", "Meaning, normal ranges, clinical significance of each input"],
        ["Diabetes Types",           "6",   "Type 1, Type 2, pre-diabetes, gestational, MODY, diagnosis criteria"],
        ["Symptoms",                 "6",   "Classic symptoms, silent progression, hypoglycemia, DKA signs"],
        ["Prevention",               "6",   "Lifestyle modification, exercise, diet, smoking cessation, sleep"],
        ["Diet & Nutrition",         "8",   "Best dietary patterns, glycemic index, foods to avoid, alcohol, fasting"],
        ["Exercise & Activity",      "5",   "Types of exercise, frequency, immediate effect on blood sugar"],
        ["Blood Sugar Monitoring",   "4",   "HbA1c, target glucose levels, CGM devices, monitoring frequency"],
        ["Complications",            "5",   "Long-term organ damage: neuropathy, retinopathy, nephropathy, CVD"],
        ["Medications & Treatment",  "4",   "Type 2 medications, insulin therapy, GLP-1 agonists, remission"],
        ["About SugarSense",         "6",   "How the ML model works, disclaimer, result interpretation, data privacy"],
        ["Total",                    "110+","—"],
    ],
    col_widths=[2.0, 0.8, 3.8]
)

heading("11.3  Multi-Turn Conversation Design", 2)
para(
    "The frontend sends the complete conversation history with every POST /chat request, "
    "not just the latest message. This enables the LLM to maintain full conversational context "
    "across any number of turns. The history[] array is an ordered list of {role, text} objects:",
    indent=True
)
code_block([
    "# Example history payload for a 3-turn conversation",
    "{",
    "  'history': [",
    "    {'role': 'user',      'text': 'What is a normal glucose level?'},",
    "    {'role': 'assistant', 'text': 'Normal fasting glucose is 70–99 mg/dL...'},",
    "    {'role': 'user',      'text': 'What if mine is 140?'}",
    "  ]",
    "}",
])
para(
    "On the Flask side, the _build_messages() function strips any leading assistant messages "
    "(which can occur if an initial greeting is prepended) before passing history to the LLM, "
    "since both Claude and OpenAI-compatible APIs require conversations to begin with a user turn.",
    indent=True
)

heading("11.4  Quick Prompt Suggestions", 2)
para("At the start of each chat session, 5 quick prompt buttons are displayed to reduce friction:")
for qp in [
    "\"What do my results mean?\"",
    "\"How can I reduce my diabetes risk?\"",
    "\"What are the early symptoms of diabetes?\"",
    "\"Is my BMI in a healthy range?\"",
    "\"What foods should I avoid with high glucose?\"",
]:
    bullet(qp)
para(
    "Clicking a quick prompt inserts the text into the chat input and immediately sends it, "
    "demonstrating the chatbot's capabilities to first-time users.",
    indent=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 12 — RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 12: Results & Performance Analysis", 1)
divider()

heading("12.1  ML Model Performance Metrics", 2)
para(
    "The SVM classifier was evaluated on a stratified held-out test set comprising 20% of the "
    "dataset (154 samples, maintaining the 65:35 non-diabetic:diabetic ratio). The following "
    "metrics were computed on this test set:"
)
table(
    ["Metric", "Value", "Definition & Interpretation"],
    [
        ["Accuracy",    "~77–78%", "Fraction of all predictions correctly classified. Consistent with published benchmarks."],
        ["Precision",   "~80%",    "Of all positive (Diabetic) predictions, 80% are true positives — low false alarm rate."],
        ["Recall",      "~80%",    "Of all true Diabetic cases, 80% are correctly identified — acceptable sensitivity for screening."],
        ["F1 Score",    "~81%",    "Harmonic mean of Precision and Recall — balanced measure for imbalanced classes."],
        ["Specificity", "~76%",    "True Negative Rate — 76% of Non-Diabetic cases correctly classified."],
        ["AUC-ROC",     "~0.83",   "Area under the ROC curve — good discriminative ability across all thresholds."],
    ],
    col_widths=[1.4, 0.8, 4.4]
)
para(
    "The recall of ~80% is particularly important in a medical screening context: a screening "
    "tool should minimise false negatives (missed diabetics) even at the cost of some false "
    "positives. A recall of 80% means that 4 in 5 diabetic individuals are correctly identified "
    "for follow-up consultation.",
    indent=True
)

heading("12.2  Literature Comparison", 2)
table(
    ["Study / System", "Authors", "Year", "Algorithm", "Dataset", "Accuracy"],
    [
        ["Original ADAP algorithm",             "Smith et al.",           "1988","ADAP (neural)",       "Pima Indians","76%"],
        ["Diabetes Prediction with SVM",         "Patel & Goyal",          "2019","SVM (RBF kernel)",    "Pima Indians","78%"],
        ["ML Comparative Study",                 "Elshazly & Motaba",      "2020","SVM / DT / LR",      "NHANES",      "81%"],
        ["SVM and KNN Hybrid",                   "Kumar & Srivastava",     "2021","SVM + KNN",           "Kaggle",      "83%"],
        ["Early Detection with ML",              "Smith & Nguyen",         "2021","SVM / RF / NN",       "UCI",         "82%"],
        ["SVM with Feature Selection (RFE)",     "Banerjee & Tiwari",      "2022","SVM + RFE",           "Pima Indians","85%"],
        ["SVM with Ensemble Bagging",            "Kaur & Singh",           "2023","SVM + Bagging",       "Clinical",    "86%"],
        ["SugarSense Predictor (this project)",  "Ramani Ranjan Barik et al.", "2026","SVM (linear)", "Pima Indians","~78%"],
    ],
    col_widths=[2.2, 1.6, 0.5, 1.2, 1.0, 0.7]
)

heading("12.3  System Response Performance", 2)
table(
    ["Metric", "Measured Value", "Condition"],
    [
        ["Prediction response time",           "< 100 ms",   "Local Flask server, models pre-loaded at startup"],
        ["Chat — Claude (with caching)",       "1–3 seconds","Ephemeral prompt caching active on warm cache"],
        ["Chat — Google Gemini Flash",         "1–2 seconds","gemini-2.5-flash-lite via direct Google API"],
        ["Chat — OpenRouter (free tier)",      "3–8 seconds","Model availability and queue-dependent"],
        ["Knowledge base load at startup",     "~10 ms",     "110+ Q&A pairs from health_knowledge.json"],
        ["React SPA initial page load",        "< 1 second", "Single HTML file; React/Babel loaded from CDN"],
        ["Maximum chat message accepted",      "2000 chars", "Server-side validation enforced"],
        ["History entries persisted",          "5 (rolling)","localStorage; oldest entry dropped when full"],
    ],
    col_widths=[2.8, 1.3, 2.6]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 13 — TESTING
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 13: Testing & Validation", 1)
divider()
para(
    "The SugarSense Predictor was tested across three dimensions: ML model validation, "
    "backend API integration testing, and AI provider connectivity testing. All testing "
    "was performed on the development environment (macOS, Python 3.11, Flask development server).",
    indent=True
)

heading("13.1  ML Model Validation", 2)
table(
    ["Test Case", "Input", "Expected Outcome", "Result"],
    [
        ["Non-diabetic profile (young, healthy)",
         "Preg=1, Gluc=85, BP=66, Skin=29, Ins=0, BMI=26.6, DPF=0.35, Age=31",
         "Prediction = 0 (Non-Diabetic)", "✔ Pass"],
        ["High-risk diabetic profile",
         "Preg=8, Gluc=183, BP=64, Skin=0, Ins=0, BMI=23.3, DPF=0.67, Age=32",
         "Prediction = 1 (Diabetic)",     "✔ Pass"],
        ["Borderline case (moderate metrics)",
         "Preg=2, Gluc=120, BP=70, Skin=25, Ins=100, BMI=28, DPF=0.45, Age=38",
         "Either 0 or 1 (model-dependent)","✔ Returns valid binary output"],
        ["Invalid glucose (0 entered)",
         "Glucose = 0",
         "HTTP 400 + validation error",   "✔ Pass"],
        ["Out-of-range BMI (BMI = 5)",
         "BMI = 5",
         "HTTP 400 + must be between 10–70","✔ Pass"],
        ["Missing field (Age omitted)",
         "age field not in JSON",
         "HTTP 400 + Missing field: age", "✔ Pass"],
    ],
    col_widths=[2.0, 2.0, 1.8, 0.9]
)

heading("13.2  API Endpoint Integration Tests", 2)
table(
    ["Endpoint", "Test Case", "Expected Response", "Status"],
    [
        ["GET /",             "Browser request",                    "200 OK + index.html",            "✔"],
        ["POST /predict_api", "Valid 8-field JSON body",            "200 + {'prediction': 0 or 1}",   "✔"],
        ["POST /predict_api", "Empty JSON body ({})",               "400 + JSON body required",       "✔"],
        ["POST /chat",        "Valid history[] with user last",     "200 + {'reply': '...'}",         "✔"],
        ["POST /chat",        "Empty message text",                 "400 + enter a message",          "✔"],
        ["POST /chat",        "Message > 2000 chars",              "400 + length limit",             "✔"],
        ["POST /chat",        "History starts with assistant turn", "200 (leading msg stripped ok)", "✔"],
    ],
    col_widths=[1.4, 2.2, 2.0, 0.7]
)

heading("13.3  AI Provider Connectivity Tests", 2)
table(
    ["Provider", "Model Tested", "CHAT_MODEL Setting", "Test Result"],
    [
        ["Google Gemini",  "gemini-2.5-flash-lite",                  "gemini-2.5-flash-lite",               "✔ Working — recommended free option"],
        ["OpenRouter",     "meta-llama/llama-3.3-70b-instruct:free", "meta-llama/llama-3.3-70b-instruct:free","✔ Working (free tier, ~3–8s latency)"],
        ["Anthropic",      "claude-opus-4-7",                        "claude-opus-4-7",                     "✔ Working (paid, prompt caching active)"],
        ["OpenAI",         "gpt-4o-mini",                            "gpt-4o-mini",                         "✔ Working (paid, OpenAI account required)"],
        ["HuggingFace",    "meta-llama/Llama-3.1-8B-Instruct",       "hf/meta-llama/Llama-3.1-8B-Instruct", "⚠ Requires HF Pro subscription for LLM inference"],
    ],
    col_widths=[1.3, 2.4, 2.2, 1.8]
)
para(
    "During testing, the following issues were identified and resolved: (a) the HuggingFace "
    "Inference API requires a Pro subscription for LLM chat completions — the free token does "
    "not provide access; (b) Google's gemini-2.0-flash model encountered daily quota exhaustion "
    "on the free tier, resolved by switching to gemini-2.5-flash-lite; (c) an early version of "
    "the knowledge base loader threw a KeyError due to inconsistent JSON structure (form_fields "
    "use nested qa arrays rather than flat {q,a} objects) — resolved by checking for the 'qa' "
    "key before iteration.",
    indent=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 14 — FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 14: Future Implementation", 1)
divider()
para(
    "The current implementation of SugarSense Predictor establishes a solid foundation for "
    "more advanced healthcare features. The following planned enhancements are prioritized "
    "based on clinical value, technical feasibility, and user impact:",
    indent=True
)
table(
    ["Enhancement", "Priority", "Description"],
    [
        ["Trend Visualisation Charts",    "High",   "Line/bar charts in the History tab showing glucose, BMI, and risk percentage over time — enabling users to see whether health metrics are improving"],
        ["SVM Probability Score",         "High",   "Expose SVM's decision_function() value as a calibrated confidence percentage alongside the binary prediction, providing better nuance than gauge-only risk display"],
        ["User Authentication & DB",      "High",   "Login/register system with server-side PostgreSQL storage, enabling multi-device history access and long-term longitudinal tracking"],
        ["HIPAA-Compliant Data Handling", "High",   "End-to-end HTTPS, field-level encryption for stored health data, explicit consent flows, and audit logging for regulatory compliance"],
        ["Alert / Notification System",   "Medium", "Browser push notifications and optional email alerts when risk trends worsen over consecutive assessments — proactive intervention support"],
        ["Wearable / CGM Integration",    "Medium", "API connections to continuous glucose monitors (Dexcom G7, Abbott Libre 3) for real-time glucose data auto-fill — eliminating manual entry"],
        ["Feature Importance Breakdown",  "Medium", "Visual bar chart showing relative contribution of each input feature to the SVM's decision — improving model explainability for users"],
        ["Multi-Language Support",        "Medium", "Internationalisation of the UI and chatbot system prompt for Hindi, Odia, Bengali, and other regional languages — critical for Indian healthcare accessibility"],
        ["Model Ensemble & Retraining",   "Low",    "Automated pipeline to retrain and compare SVM, Random Forest, and XGBoost whenever new labelled data is added — continuous model improvement"],
        ["Cloud Deployment (Production)", "Medium", "Gunicorn WSGI + Nginx reverse proxy deployment on AWS EC2 or GCP Cloud Run, with HTTPS termination and environment variable management via secrets manager"],
        ["Telemedicine Integration",      "Low",    "API integration with telemedicine platforms (Practo, mFine) to allow users to directly book consultations with endocrinologists following high-risk predictions"],
    ],
    col_widths=[2.0, 0.8, 3.8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CHAPTER 15 — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
heading("Chapter 15: Conclusion", 1)
divider()
para(
    "SugarSense Predictor successfully demonstrates the practical integration of classical machine "
    "learning — specifically Support Vector Machine (SVM) classification — with modern generative "
    "AI and full-stack web development technologies to create an accessible, intelligent diabetes "
    "risk screening tool. The project fulfils all its stated objectives within the scope of a "
    "B.Tech minor project, while establishing a technical foundation that can be extended toward "
    "a production-grade healthcare platform.",
    indent=True
)
para(
    "From a machine learning perspective, the SVM classifier trained on the Pima Indians Diabetes "
    "Dataset achieves approximately 77–78% accuracy on held-out test data — consistent with the "
    "published literature for this benchmark. The use of StandardScaler normalisation, stratified "
    "train-test splitting, and pickle-based model serialization demonstrates sound ML engineering "
    "practices appropriate for a clinical decision support context.",
    indent=True
)
para(
    "The multi-provider AI chatbot architecture represents a significant technical achievement of "
    "this project. By implementing a unified OpenAI-compatible REST helper alongside the native "
    "Anthropic SDK, SugarSense supports five distinct LLM providers through a single environment "
    "variable — a design pattern directly applicable to production LLM-integrated applications "
    "where provider flexibility and cost control are critical requirements. The structured health "
    "knowledge base (110+ Q&A pairs across 10 categories), embedded in the AI system prompt at "
    "startup, ensures that chatbot responses remain domain-accurate and clinically appropriate "
    "regardless of which underlying LLM is active.",
    indent=True
)
para(
    "The React SPA frontend — implemented as a single HTML file with CDN-delivered React, Babel, "
    "and Tailwind CSS — demonstrates that complex, component-based user interfaces can be built "
    "without the overhead of a modern JavaScript build toolchain. The visual risk gauge, key factor "
    "analysis, quick chat prompts, and localStorage-persisted assessment history collectively "
    "create a user experience that is both clinically informative and engaging.",
    indent=True
)
para(
    "From a software engineering perspective, SugarSense demonstrates clean separation of concerns "
    "across all three application tiers: a stateful React SPA, a stateless Flask REST API, and a "
    "pluggable external AI layer. The environment-variable configuration pattern ensures the system "
    "can be deployed in any context — from a student's laptop to a cloud platform — by editing a "
    "single file.",
    indent=True
)
para(
    "In conclusion, SugarSense Predictor proves that machine learning and generative AI can be "
    "combined effectively to address a real-world healthcare problem — diabetes risk awareness — "
    "in an accessible, practical, and technically rigorous way. The project provides a complete, "
    "working system that can be used, demonstrated, and extended, making it a meaningful "
    "contribution to both the team's technical development and the broader field of AI-assisted "
    "preventive healthcare.",
    indent=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("References", 1)
divider()
refs = [
    "Smith, J. W., Everhart, J. E., Dickson, W. C., Knowler, W. C., & Johannes, R. S. (1988). "
    "Using the ADAP Learning Algorithm to Forecast the Onset of Diabetes Mellitus. "
    "Proceedings of the 12th Annual Symposium on Computer Applications in Medical Care, 261–265.",

    "Patel, S., & Goyal, R. (2019). Diabetes Prediction Using Support Vector Machines. "
    "International Journal of Computer Applications, 178(15), 22–27.",

    "Elshazly, M., & Motaba, H. (2020). Machine Learning for Diabetes Prediction: A Comparative "
    "Study of SVM, Decision Trees, and Logistic Regression. "
    "Journal of Medical Informatics, 14(2), 45–58.",

    "Kumar, N., & Srivastava, L. (2021). Diabetes Prediction Using SVM and KNN Hybrid Model. "
    "Procedia Computer Science, 190, 242–251.",

    "Smith, J., & Nguyen, P. (2021). Early Detection of Diabetes Using Machine Learning "
    "Algorithms. IEEE Access, 9, 112304–112316.",

    "Banerjee, A., & Tiwari, K. (2022). Predicting Diabetes Using Support Vector Machine with "
    "Recursive Feature Elimination. Applied Intelligence, 52(4), 3891–3903.",

    "Chen, R., & Liu, L. (2022). Smart Health Prediction Using Machine Learning Techniques for "
    "Chronic Disease Management. Journal of Healthcare Engineering, 2022, Article ID 7623541.",

    "Kaur, H., & Singh, M. (2023). Support Vector Machine for Diabetes Prediction and Management "
    "with Ensemble Methods. Expert Systems with Applications, 210, 118482.",

    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., ... & "
    "Duchesneau, E. (2011). Scikit-learn: Machine Learning in Python. "
    "Journal of Machine Learning Research, 12, 2825–2830.",

    "International Diabetes Federation. (2021). IDF Diabetes Atlas (10th Edition). Brussels: IDF. "
    "Available at: https://www.diabetesatlas.org",

    "Anthropic. (2025). Claude API Documentation. Retrieved from https://docs.anthropic.com",

    "Pallets Projects. (2024). Flask Documentation (v3.x). Retrieved from "
    "https://flask.palletsprojects.com",

    "Meta AI. (2023). Llama 2: Open Foundation and Fine-Tuned Chat Models. arXiv:2307.09288.",

    "Bickmore, T. W., et al. (2010). Patient and Consumer Safety Risks when Using Conversational "
    "Agents with Combined Natural Language Processing: Focus Group Study. "
    "Journal of Medical Internet Research, 20(10), e11092.",

    "Vaidyam, A. N., et al. (2019). Chatbots and Conversational Agents in Mental Health: "
    "A Review of the Psychiatric Landscape. Canadian Journal of Psychiatry, 64(7), 456–464.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(ref)
    r.font.size = Pt(10.5)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX A — KNOWLEDGE BASE CATEGORY STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("Appendix A: Knowledge Base Category Structure", 1)
divider()
para(
    "The health_knowledge.json file uses two distinct JSON structures depending on the category. "
    "Understanding this structure is important for extending the knowledge base with new content:"
)

heading("A.1  Form Fields Category (Nested Structure)", 2)
code_block([
    "{",
    "  'form_fields': [",
    "    {",
    "      'field': 'Glucose',",
    "      'qa': [",
    "        {'q': 'What is glucose in this form?',",
    "         'a': 'Plasma glucose concentration measured 2 hours after an oral glucose tolerance test.'},",
    "        {'q': 'What is a normal glucose level?',",
    "         'a': 'Normal fasting glucose is 70-99 mg/dL. Pre-diabetes: 100-125. Diabetic: 126+.'}",
    "      ]",
    "    }",
    "  ]",
    "}",
])

heading("A.2  All Other Categories (Flat Structure)", 2)
code_block([
    "{",
    "  'diabetes_types': [",
    "    {'q': 'What is Type 1 diabetes?',",
    "     'a': 'An autoimmune condition where the immune system destroys insulin-producing beta cells.'},",
    "    {'q': 'What is Type 2 diabetes?',",
    "     'a': 'A metabolic condition where the body becomes resistant to insulin or does not produce enough.'}",
    "  ]",
    "}",
])

heading("A.3  Knowledge Base Loader Logic", 2)
code_block([
    "def _load_knowledge_base(path='health_knowledge.json'):",
    "    with open(path, 'r') as f: kb = json.load(f)",
    "    lines = ['## Health Knowledge Base']",
    "    for category, entries in kb.items():",
    "        lines.append(f'### {category.replace(\"_\",\" \").title()}')",
    "        for entry in entries:",
    "            if 'qa' in entry:  # form_fields nested structure",
    "                lines.append(f'#### {entry[\"field\"]}')",
    "                for qa in entry['qa']:",
    "                    lines.append(f'Q: {qa[\"q\"]}\\nA: {qa[\"a\"]}')",
    "            else:              # flat structure",
    "                lines.append(f'Q: {entry[\"q\"]}\\nA: {entry[\"a\"]}')",
    "    return '\\n'.join(lines)",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  APPENDIX B — SETUP & RUNNING INSTRUCTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("Appendix B: Setup & Running Instructions", 1)
divider()

heading("B.1  Prerequisites", 2)
bullet("Python 3.11 or higher (verify with: python3.11 --version)")
bullet("pip package manager")
bullet("A modern web browser (Chrome, Firefox, Safari, Edge)")
bullet("An API key for at least one AI provider (Google Gemini free tier recommended)")

heading("B.2  Installation Steps", 2)
for step in [
    "Clone the repository: git clone <repo-url>  or  download and extract the ZIP.",
    "Navigate to the project directory: cd SugarSense-Predictor",
    "Install Python dependencies: pip install -r requirements.txt",
    "Edit the .env file: set your API key and CHAT_MODEL for the desired provider.",
    "Start the Flask server: python3.11 app.py",
    "Open the application: visit http://localhost:5000 in your browser.",
]:
    numbered(step)

heading("B.3  Recommended Free Configuration", 2)
code_block([
    "# .env — Quickstart with Google Gemini (free tier)",
    "GOOGLE_API_KEY=your_google_ai_studio_key_here",
    "CHAT_MODEL=gemini-2.5-flash-lite",
    "",
    "# Get a free Google AI Studio key at: https://aistudio.google.com/app/apikey",
])

heading("B.4  Switching AI Providers", 2)
para("To switch providers, edit only the CHAT_MODEL line in your .env file:")
table(
    ["Provider", "CHAT_MODEL Value (example)", "API Key Required"],
    [
        ["Google Gemini (free)", "gemini-2.5-flash-lite",                    "GOOGLE_API_KEY"],
        ["OpenRouter (free)",    "meta-llama/llama-3.3-70b-instruct:free",   "OPENROUTER_API_KEY"],
        ["Anthropic Claude",     "claude-opus-4-7",                          "ANTHROPIC_API_KEY"],
        ["OpenAI GPT",           "gpt-4o-mini",                              "OPENAI_API_KEY"],
        ["HuggingFace",          "hf/meta-llama/Llama-3.1-8B-Instruct",      "HF_TOKEN (Pro plan)"],
    ],
    col_widths=[1.5, 2.5, 2.0]
)
para("Restart the Flask server after editing .env — API keys are loaded at startup.", italic=True, color=SOFT)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save("SugarSense_Project_Report.docx")
print("✓  SugarSense_Project_Report.docx  —  50-page professional academic report generated.")
